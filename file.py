import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime
from fpdf import FPDF
import base64

st.set_page_config(page_title="AI Resume Builder Ultra", page_icon="🧠", layout="wide")

# Custom Theme and Styling
st.markdown("""
    <style>
        body {
            background-color: #f0f2f6;
        }
        .block-container {
            padding-top: 2rem;
        }
        .stTextInput > div > div > input,
        .stTextArea > div > textarea {
            border-radius: 12px;
            border: 1px solid #ccc;
            padding: 0.75rem;
        }
        .stButton>button {
            background: linear-gradient(to right, #00c6ff, #0072ff);
            color: white;
            font-weight: bold;
            padding: 10px 20px;
            border-radius: 10px;
        }
    </style>
""", unsafe_allow_html=True)

st.title("🧠 AI Resume Builder Ultra")
st.markdown("#### ✨ Create an impressive, professional resume in minutes using smart automation and beautiful formatting.")

# Personal Section
st.sidebar.header("👤 Personal Info")
with st.sidebar:
    name = st.text_input("Full Name", "Kajola Gbenga")
    email = st.text_input("Email", "k.gbenga234@gmail.com")
    phone = st.text_input("Phone", "09038780790")
    location = st.text_input("Location", "Oyo, Nigeria")
    linkedin = st.text_input("LinkedIn URL", "https://linkedin.com")

# Main Form Sections
st.markdown("---")
st.subheader("📝 Professional Summary")
summary = st.text_area("Write a short professional summary", "An awesome data analyst with an exceptional track record.", height=120)

st.subheader("💡 Core Skills")
cols = st.columns(2)
with cols[0]:
    prog_skills = st.text_input("Programming Languages (comma-separated)", "Python, JavaScript, SQL").split(',')
    viz_tools = st.text_input("Visualization Tools", "Seaborn, Matplotlib").split(',')
with cols[1]:
    ml_skills = st.text_input("ML & AI Libraries", "Tensorflow, Random Forest").split(',')
    db_tools = st.text_input("Databases & Tools", "MySQL, PostgreSQL").split(',')

st.markdown("---")
st.subheader("💼 Work Experience")
exp_count = st.slider("How many jobs have you had?", 1, 5, 2)
experiences = []
for i in range(exp_count):
    with st.expander(f"Experience #{i+1}"):
        title = st.text_input(f"Job Title {i+1}", f"Data Analyst", key=f"title{i}")
        company = st.text_input(f"Company {i+1}", "Freelance", key=f"company{i}")
        dates = st.text_input(f"Dates {i+1}", "Feb 2021 - Aug 2022", key=f"dates{i}")
        details = st.text_area(f"Achievements {i+1} (comma-separated)", "Built dashboards, Analyzed trends", key=f"details{i}").split(',')
        experiences.append({"title": title, "company": company, "dates": dates, "details": details})

st.markdown("---")
st.subheader("🎓 Education")
edu = st.text_area("Your Education Info", "BSc, Olabisi Onabanjo University, 2011 - 2015", height=80)

st.subheader("📜 Certifications")
certs = st.text_area("List your Certifications (comma-separated)", "Google Data Analytics, IBM Data Science").split(',')

st.subheader("🌍 Leadership & Community Engagement")
leadership = st.text_area("Mention any Leadership or Volunteer roles (comma-separated)", "Data mentor, Open source contributor").split(',')

# Resume Generation Logic
def generate_resume():
    doc = Document()
    doc.add_heading(name, 0)
    doc.add_paragraph(f"Email: {email} | Phone: {phone} | Location: {location} | LinkedIn: {linkedin}")
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(f"Programming: {', '.join([s.strip() for s in prog_skills])}")
    doc.add_paragraph(f"ML & AI: {', '.join([s.strip() for s in ml_skills])}")
    doc.add_paragraph(f"Visualization: {', '.join([s.strip() for s in viz_tools])}")
    doc.add_paragraph(f"Databases & Tools: {', '.join([s.strip() for s in db_tools])}")

    doc.add_heading("Work Experience", level=1)
    for exp in experiences:
        doc.add_paragraph(f"{exp['title']} at {exp['company']} ({exp['dates']})", style='List Bullet')
        for item in exp['details']:
            if item.strip():
                doc.add_paragraph(f"- {item.strip()}", style='List Bullet 2')

    doc.add_heading("Education", level=1)
    doc.add_paragraph(edu)

    doc.add_heading("Certifications", level=1)
    for c in certs:
        if c.strip():
            doc.add_paragraph(c.strip(), style='List Bullet')

    doc.add_heading("Leadership & Community Engagement", level=1)
    for l in leadership:
        if l.strip():
            doc.add_paragraph(l.strip(), style='List Bullet')

    return doc

def generate_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt=name, ln=1, align='C')
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, txt=f"Email: {email}\nPhone: {phone}\nLocation: {location}\nLinkedIn: {linkedin}")
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Professional Summary", ln=1)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, txt=summary)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Skills", ln=1)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, txt=f"Programming: {', '.join(prog_skills)}")
    pdf.multi_cell(0, 10, txt=f"ML & AI: {', '.join(ml_skills)}")
    pdf.multi_cell(0, 10, txt=f"Visualization: {', '.join(viz_tools)}")
    pdf.multi_cell(0, 10, txt=f"Databases & Tools: {', '.join(db_tools)}")

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Work Experience", ln=1)
    pdf.set_font("Arial", '', 12)
    for exp in experiences:
        pdf.multi_cell(0, 10, txt=f"{exp['title']} at {exp['company']} ({exp['dates']})")
        for item in exp['details']:
            if item.strip():
                pdf.multi_cell(0, 10, txt=f"  - {item.strip()}")

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Education", ln=1)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, txt=edu)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Certifications", ln=1)
    pdf.set_font("Arial", '', 12)
    for c in certs:
        if c.strip():
            pdf.multi_cell(0, 10, txt=c.strip())

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Leadership & Community Engagement", ln=1)
    pdf.set_font("Arial", '', 12)
    for l in leadership:
        if l.strip():
            pdf.multi_cell(0, 10, txt=l.strip())

    # Write to BytesIO using output(dest='S')
    pdf_output = BytesIO()
    pdf_bytes = pdf.output(dest='S').encode('latin-1')
    pdf_output.write(pdf_bytes)
    pdf_output.seek(0)
    return pdf_output

if st.button("🚀 Generate My Resume"):
    final_doc = generate_resume()
    docx_buffer = BytesIO()
    final_doc.save(docx_buffer)
    docx_buffer.seek(0)
    pdf_buffer = generate_pdf()
    st.success("🎉 Your resume is ready!")
    st.download_button("⬇️ Download .docx Resume", docx_buffer, file_name=f"{name.replace(' ', '_')}_Resume.docx")
    st.download_button("⬇️ Download .pdf Resume", pdf_buffer, file_name=f"{name.replace(' ', '_')}_Resume.pdf")

# Footer
st.markdown("---")
st.markdown("# About the Developer")
# Display developer image
st.image("My image6.jpg", width=250)
st.markdown("## **Kajola Gbenga**")

st.markdown(
    """
📇 Certified Data Analyst | Certified Data Scientist | Certified SQL Programmer | Mobile App Developer | AI/ML Engineer

🔗 [LinkedIn](https://www.linkedin.com/in/kajolagbenga)  
📜 [View My Certifications & Licences](https://www.datacamp.com/portfolio/kgbenga234)  
💻 [GitHub](https://github.com/prodigy234)  
🌐 [Portfolio](https://kajolagbenga.netlify.app/)  
📧 k.gbenga234@gmail.com
"""
)


st.markdown("✅ Created using Python and Streamlit")